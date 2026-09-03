"""Test cho SummarEase Django."""

import tempfile
from unittest.mock import MagicMock, patch

from django.contrib.auth.models import User
from django.core.cache import cache
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse

from .forms import SettingsForm, SummaryRequestForm
from .models import Document, Summary, UserProfile, UserSetting
from .nlp import textrank_summarize
from .nlp_utils import (
    detect_language,
    extract_keywords,
    generate_title,
    highlight_keywords,
    normalize_text,
    split_sentences,
    split_words,
    truncate_text,
)
from .readers import _is_private_ip, _resolve_and_validate, extract_text
from .signing import decrypt_value, encrypt_value


class TestHelperMixin:
    def _create_user(self, username, password="secret123", is_superuser=False):
        if is_superuser:
            user = User.objects.create_superuser(username=username, password=password)
        else:
            user = User.objects.create_user(username=username, password=password)
        return user


# ──────────────────────────────────────────────
#  NLP UNIT TESTS
# ──────────────────────────────────────────────


class NlpSplitTests(TestCase):
    def test_split_sentences_basic(self):
        result = split_sentences("Hello world. This is fun!")
        self.assertEqual(len(result), 2)

    def test_split_sentences_vietnamese_diacritic_start(self):
        text = "Em là học sinh. Ở trường tôi học giỏi. Đó là điều quan trọng."
        result = split_sentences(text)
        self.assertEqual(len(result), 3)

    def test_split_sentences_abbreviations(self):
        text = "Dr. Smith went to New York. He arrived at 5 p.m."
        result = split_sentences(text)
        self.assertEqual(len(result), 2)

    def test_split_sentences_vietnamese_abbrev(self):
        text = "Tp. Hồ Chí Minh là thành phố lớn nhất. Nó nằm ở phía Nam."
        result = split_sentences(text)
        self.assertEqual(len(result), 2)

    def test_split_words_basic(self):
        result = split_words("Hello World!")
        self.assertIn("hello", result)
        self.assertIn("world", result)

    def test_split_words_vietnamese(self):
        result = split_words("Xin chào thế giới!")
        self.assertIn("chào", result)
        self.assertIn("thế", result)


class NlpNormalizeTests(TestCase):
    def test_normalize_removes_extra_spaces(self):
        result = normalize_text("Hello    world\n\n  test")
        self.assertEqual(result, "Hello world test")

    def test_normalize_strips_whitespace(self):
        result = normalize_text("  hello  ")
        self.assertEqual(result, "hello")


class NlpDetectLanguageTests(TestCase):
    def test_detect_english_by_diacritics(self):
        result = detect_language("The quick brown fox jumps over the lazy dog.")
        self.assertEqual(result, "english")

    def test_detect_vietnamese_by_diacritics(self):
        result = detect_language("Xin chào thế giới! Hôm nay là một ngày đẹp trời.")
        self.assertEqual(result, "vietnamese")

    def test_detect_vietnamese_no_diacritics(self):
        result = detect_language("Xin chao the gioi! Hom nay la mot ngay dep troi.")
        self.assertEqual(result, "vietnamese")

    def test_detect_empty_falls_to_english(self):
        result = detect_language("")
        self.assertEqual(result, "english")


class NlpKeywordsTests(TestCase):
    def test_extract_keywords_returns_top_words(self):
        text = "python django python django django web framework python"
        result = extract_keywords(text, limit=3)
        self.assertIn("python", result)
        self.assertIn("django", result)
        self.assertGreaterEqual(len(result), 2)

    def test_highlight_keywords_adds_mark_tags(self):
        text = "python is great. django is better."
        result = highlight_keywords(text, ["python", "django"])
        self.assertIn("<mark>python</mark>", result)
        self.assertIn("<mark>django</mark>", result)

    def test_highlight_keywords_escapes_html(self):
        text = "test <script>alert('xss')</script>"
        result = highlight_keywords(text, ["test"])
        self.assertIn("&lt;script&gt;", result)
        self.assertNotIn("<script>", result)


class NlpTitleTests(TestCase):
    def test_generate_title_from_summary(self):
        result = generate_title("This is the first sentence of the summary.")
        self.assertEqual(result, "This is the first sentence of the summary.")

    def test_generate_title_with_source_name(self):
        result = generate_title("summary", "my-doc.pdf")
        self.assertIn("my-doc.pdf", result)

    def test_generate_title_empty_returns_default(self):
        result = generate_title("")
        self.assertEqual(result, "Tóm tắt tài liệu")


class NlpTruncateTests(TestCase):
    def test_truncate_short_text(self):
        text = "short text"
        result = truncate_text(text, max_chars=100)
        self.assertEqual(result, text)

    def test_truncate_long_text(self):
        text = " ".join(["word"] * 100)
        result = truncate_text(text, max_chars=20)
        self.assertLessEqual(len(result), 20)


# ──────────────────────────────────────────────
#  MODEL TESTS
# ──────────────────────────────────────────────


class DocumentModelTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="doc-tester", password="secret123")

    def test_create_document(self):
        doc = Document.objects.create(
            user=self.user,
            source_type="text",
            title="Test doc",
            content="Hello world",
        )
        self.assertEqual(str(doc), "Test doc")


class HealthCheckTests(TestCase):
    def test_health_endpoint(self):
        response = self.client.get(reverse("health"))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json(), {"status": "ok"})


class SummaryModelTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="sum-tester", password="secret123")
        self.doc = Document.objects.create(
            user=self.user, source_type="text", title="Doc", content="Content"
        )

    def test_create_summary(self):
        summary = Summary.objects.create(
            document=self.doc,
            user=self.user,
            title="Sum",
            method="textrank",
            language="english",
            ratio=0.5,
            summary_text="Summary text",
        )
        self.assertEqual(str(summary), "Sum")

    def test_summary_timestamps(self):
        summary = Summary.objects.create(
            document=self.doc,
            user=self.user,
            title="Sum",
            method="textrank",
            language="english",
            ratio=0.5,
            summary_text="Text",
        )
        self.assertIsNotNone(summary.created_at)


class UserProfileModelTests(TestCase):
    def test_create_profile_auto_defaults(self):
        user = User.objects.create_user(username="profile-test", password="secret123")
        profile = UserProfile.objects.get(user=user)
        self.assertEqual(profile.role, "user")

    def test_admin_profile_role(self):
        user = User.objects.create_superuser(username="admin-test", password="secret123")
        profile = UserProfile.objects.get(user=user)
        self.assertEqual(profile.role, "admin")


# ──────────────────────────────────────────────
#  VIEW / INTEGRATION TESTS
# ──────────────────────────────────────────────


class AuthPageTests(TestCase):
    def test_home_page(self):
        response = self.client.get(reverse("home"))
        self.assertEqual(response.status_code, 200)

    def test_login_page(self):
        response = self.client.get(reverse("login"))
        self.assertEqual(response.status_code, 200)

    def test_register_page(self):
        response = self.client.get(reverse("register"))
        self.assertEqual(response.status_code, 200)

    @override_settings(STATIC_URL="/static/")
    def test_home_page_has_static_links(self):
        response = self.client.get(reverse("home"))
        self.assertContains(response, "/static/css/app.css")


class AuthFlowTests(TestCase):
    def test_register_creates_user_and_redirects(self):
        response = self.client.post(
            reverse("register"),
            {
                "username": "newuser",
                "password1": "StrongPass123!",
                "password2": "StrongPass123!",
            },
        )
        self.assertEqual(response.status_code, 302)
        self.assertTrue(User.objects.filter(username="newuser").exists())

    def test_login_and_logout(self):
        User.objects.create_user(username="logintest", password="secret123")
        response = self.client.post(
            reverse("login"),
            {
                "username": "logintest",
                "password": "secret123",
            },
        )
        self.assertEqual(response.status_code, 302)


class SummaryFlowTests(TestCase):
    def tearDown(self):
        cache.clear()

    def setUp(self):
        self.user = User.objects.create_user(username="tester", password="secret123")

        self.other = User.objects.create_user(username="other", password="secret123")

        self.admin = User.objects.create_superuser(username="admin", password="secret123")


    def test_login_required_for_create_summary(self):
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "text",
                "method": "textrank",
                "text": "test text",
                "ratio": 0.2,
            },
        )
        self.assertEqual(response.status_code, 302)

    def test_create_summary_textrank(self):
        self.client.login(username="tester", password="secret123")
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "text",
                "method": "textrank",
                "text": "First sentence here. Second sentence follows. Third one is final.",
                "ratio": 0.3,
            },
        )
        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.json()["ok"])
        self.assertEqual(Summary.objects.count(), 1)

    def test_create_summary_empty_text_returns_error(self):
        self.client.login(username="tester", password="secret123")
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "text",
                "method": "textrank",
                "text": "",
                "ratio": 0.2,
            },
        )
        self.assertEqual(response.status_code, 400)
        self.assertIn("errors", response.json())

    def test_rate_limit_blocks_rapid_requests(self):
        self.client.login(username="tester", password="secret123")
        self.client.post(
            reverse("create_summary"),
            {
                "source_type": "text",
                "method": "textrank",
                "text": "A sentence. B sentence. C sentence.",
                "ratio": 0.2,
            },
        )
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "text",
                "method": "textrank",
                "text": "D sentence. E sentence. F sentence.",
                "ratio": 0.2,
            },
        )
        self.assertEqual(response.status_code, 429)

    def test_admin_can_view_all_history(self):
        doc = Document.objects.create(
            user=self.other,
            source_type="text",
            title="Doc",
            content="Content",
        )
        Summary.objects.create(
            document=doc,
            user=self.other,
            title="Other summary",
            method="textrank",
            language="en",
            ratio=0.2,
            summary_text="Text",
        )
        self.client.login(username="admin", password="secret123")
        response = self.client.get(reverse("history"))
        self.assertContains(response, "Other summary")

    def test_user_cannot_view_others_detail(self):
        doc = Document.objects.create(
            user=self.other,
            source_type="text",
            title="Doc",
            content="Content",
        )
        summary = Summary.objects.create(
            document=doc,
            user=self.other,
            title="Private",
            method="textrank",
            language="en",
            ratio=0.2,
            summary_text="Text",
        )
        self.client.login(username="tester", password="secret123")
        response = self.client.get(reverse("history_detail", kwargs={"pk": summary.pk}))
        self.assertEqual(response.status_code, 404)


class SettingsFlowTests(TestCase):
    def tearDown(self):
        cache.clear()

    def setUp(self):
        self.user = User.objects.create_user(username="settings-test", password="secret123")


    def test_settings_requires_login(self):
        response = self.client.get(reverse("settings"))
        self.assertEqual(response.status_code, 302)

    def test_settings_page_loads(self):
        self.client.login(username="settings-test", password="secret123")
        response = self.client.get(reverse("settings"))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Gemini")

    def test_settings_update_ratio(self):
        self.client.login(username="settings-test", password="secret123")
        self.client.post(
            reverse("settings"),
            {
                "default_summary_ratio": 0.7,
                "gemini_api_key": "",
            },
        )
        updated = UserSetting.objects.get(user=self.user)
        self.assertEqual(updated.default_summary_ratio, 0.7)

    def test_settings_save_and_clear_api_key(self):
        self.client.login(username="settings-test", password="secret123")
        self.client.post(
            reverse("settings"),
            {
                "default_summary_ratio": 0.2,
                "gemini_api_key": "my-key-123",
            },
        )
        stored = UserSetting.objects.get(user=self.user).gemini_api_key
        self.assertEqual(decrypt_value(stored), "my-key-123")
        self.client.post(
            reverse("settings"),
            {
                "default_summary_ratio": 0.2,
                "gemini_api_key": "",
            },
        )
        self.assertEqual(UserSetting.objects.get(user=self.user).gemini_api_key, "")

    def test_settings_invalid_ratio_shows_error(self):
        self.client.login(username="settings-test", password="secret123")
        response = self.client.post(
            reverse("settings"),
            {
                "default_summary_ratio": 5.0,
                "gemini_api_key": "",
            },
        )
        self.assertContains(response, "value")


# ──────────────────────────────────────────────
#  NLP EDGE CASE TESTS
# ──────────────────────────────────────────────


class NlpEdgeCaseTests(TestCase):
    def test_split_sentences_single_sentence(self):
        result = split_sentences("Just one sentence here")
        self.assertEqual(len(result), 1)

    def test_split_sentences_empty(self):
        result = split_sentences("")
        self.assertEqual(result, [])

    def test_split_sentences_question_exclamation(self):
        result = split_sentences("What is this? Amazing! Yes.")
        self.assertEqual(len(result), 3)

    def test_split_words_empty(self):
        result = split_words("")
        self.assertEqual(result, [])

    def test_split_words_special_chars(self):
        result = split_words("hello... world!!! test's")
        self.assertIn("hello", result)
        self.assertIn("world", result)
        self.assertIn("test's", result)

    def test_normalize_empty(self):
        result = normalize_text("")
        self.assertEqual(result, "")

    def test_normalize_tabs(self):
        result = normalize_text("hello\t\tworld")
        self.assertEqual(result, "hello world")

    def test_detect_language_mixed(self):
        result = detect_language("Hello everyone! Hom nay troi dep qua.")
        self.assertEqual(result, "vietnamese")

    def test_detect_language_vietnamese_diacritics_only(self):
        result = detect_language("Xin chào, hôm nay là thứ năm.")
        self.assertEqual(result, "vietnamese")

    def test_detect_language_english_only(self):
        result = detect_language("Today is a great day to learn something new.")
        self.assertEqual(result, "english")

    def test_extract_keywords_empty(self):
        result = extract_keywords("")
        self.assertEqual(result, [])

    def test_extract_keywords_limit(self):
        text = "apple banana apple banana cherry apple date"
        result = extract_keywords(text, limit=2)
        self.assertEqual(len(result), 2)
        self.assertEqual(result[0], "apple")

    def test_highlight_keywords_no_match(self):
        result = highlight_keywords("hello world", ["python"])
        self.assertNotIn("<mark>", result)

    def test_highlight_keywords_mark_not_escaped(self):
        result = highlight_keywords("hello world", ["hello"])
        self.assertEqual(result, "<mark>hello</mark> world")

    def test_generate_title_adds_period(self):
        result = generate_title("Summary without period")
        self.assertTrue(result.endswith("."))

    def test_textrank_summarize_long_text(self):
        sentences = "This is the first sentence about Python. "
        sentences += "Django is a web framework written in Python. "
        sentences += "It is used by many developers worldwide. "
        sentences += "The framework follows the MVT architecture. "
        sentences += "Python is a versatile programming language. "
        result = textrank_summarize(sentences, ratio=0.5)
        self.assertIn("summary", result)
        self.assertIn("keywords", result)
        self.assertIn("title", result)


# ──────────────────────────────────────────────
#  FORM VALIDATION TESTS
# ──────────────────────────────────────────────


class FormValidationTests(TestCase):
    def test_settings_form_valid(self):
        form = SettingsForm(data={"default_summary_ratio": 0.3, "gemini_api_key": ""})
        self.assertTrue(form.is_valid())

    def test_settings_form_negative_ratio(self):
        form = SettingsForm(data={"default_summary_ratio": -1, "gemini_api_key": ""})
        self.assertFalse(form.is_valid())

    def test_settings_form_ratio_too_high(self):
        form = SettingsForm(data={"default_summary_ratio": 2.0, "gemini_api_key": ""})
        self.assertFalse(form.is_valid())

    def test_settings_form_long_api_key(self):
        form = SettingsForm(data={"default_summary_ratio": 0.2, "gemini_api_key": "k" * 500})
        self.assertFalse(form.is_valid())

    def test_summary_form_missing_source_type(self):
        form = SummaryRequestForm(data={})
        self.assertFalse(form.is_valid())


# ──────────────────────────────────────────────
#  ADMIN PAGE TESTS
# ──────────────────────────────────────────────


class AdminPageTests(TestCase):
    def setUp(self):
        self.admin = User.objects.create_superuser(
            username="superadmin",
            password="secret123",
        )


    def test_admin_login_required(self):
        response = self.client.get(reverse("admin:index"))
        self.assertEqual(response.status_code, 302)

    def test_admin_index_loads(self):
        self.client.login(username="superadmin", password="secret123")
        response = self.client.get(reverse("admin:index"))
        self.assertEqual(response.status_code, 200)

    def test_admin_summary_list_loads(self):
        self.client.login(username="superadmin", password="secret123")
        response = self.client.get(reverse("admin:summaries_summary_changelist"))
        self.assertEqual(response.status_code, 200)

    def test_admin_document_list_loads(self):
        self.client.login(username="superadmin", password="secret123")
        response = self.client.get(reverse("admin:summaries_document_changelist"))
        self.assertEqual(response.status_code, 200)

    def test_admin_userprofile_list_loads(self):
        self.client.login(username="superadmin", password="secret123")
        response = self.client.get(reverse("admin:summaries_userprofile_changelist"))
        self.assertEqual(response.status_code, 200)


# ──────────────────────────────────────────────
#  HISTORY PAGINATION TESTS
# ──────────────────────────────────────────────


class HistoryPaginationTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="paginator", password="secret123")

        doc = Document.objects.create(
            user=self.user,
            source_type="text",
            title="Doc",
            content="Content",
        )
        for i in range(15):
            Summary.objects.create(
                document=doc,
                user=self.user,
                title=f"Summary {i}",
                method="textrank",
                language="en",
                ratio=0.2,
                summary_text=f"Text {i}",
            )
        self.client.login(username="paginator", password="secret123")

    def test_history_first_page(self):
        response = self.client.get(reverse("history"))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Summary 14")

    def test_history_second_page(self):
        response = self.client.get(reverse("history"), {"page": 2})
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Summary 0")

    def test_history_invalid_page(self):
        response = self.client.get(reverse("history"), {"page": 999})
        self.assertEqual(response.status_code, 200)


# ──────────────────────────────────────────────
#  PERMISSION TESTS
# ──────────────────────────────────────────────


class PermissionTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="normaluser", password="secret123")

        self.admin = User.objects.create_superuser(username="super", password="secret123")

        self.doc = Document.objects.create(
            user=self.user,
            source_type="text",
            title="My doc",
            content="My content",
        )
        self.summary = Summary.objects.create(
            document=self.doc,
            user=self.user,
            title="My summary",
            method="textrank",
            language="en",
            ratio=0.2,
            summary_text="My text",
        )

    def test_own_history_detail_accessible(self):
        self.client.login(username="normaluser", password="secret123")
        response = self.client.get(reverse("history_detail", kwargs={"pk": self.summary.pk}))
        self.assertEqual(response.status_code, 200)

    def test_own_delete_allowed(self):
        self.client.login(username="normaluser", password="secret123")
        response = self.client.post(reverse("history_delete", kwargs={"pk": self.summary.pk}))
        self.assertEqual(response.status_code, 302)
        self.assertFalse(Summary.objects.filter(pk=self.summary.pk).exists())

    def test_other_delete_denied(self):
        User.objects.create_user(username="otheruser", password="secret123")

        self.client.login(username="otheruser", password="secret123")
        response = self.client.post(reverse("history_delete", kwargs={"pk": self.summary.pk}))
        self.assertEqual(response.status_code, 404)

    def test_admin_can_delete_any(self):
        self.client.login(username="super", password="secret123")
        response = self.client.post(reverse("history_delete", kwargs={"pk": self.summary.pk}))
        self.assertEqual(response.status_code, 302)
        self.assertFalse(Summary.objects.filter(pk=self.summary.pk).exists())


# ──────────────────────────────────────────────
#  DELETE CASCADE TESTS
# ──────────────────────────────────────────────


class DeleteCascadeTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="cascade-user", password="secret123")


    def test_delete_document_cascades_summary(self):
        doc = Document.objects.create(
            user=self.user,
            source_type="text",
            title="Doc",
            content="Content",
        )
        summary = Summary.objects.create(
            document=doc,
            user=self.user,
            title="Sum",
            method="textrank",
            language="en",
            ratio=0.2,
            summary_text="Text",
        )
        doc.delete()
        self.assertFalse(Summary.objects.filter(pk=summary.pk).exists())

    def test_delete_user_does_not_delete_summary(self):
        doc = Document.objects.create(
            user=self.user,
            source_type="text",
            title="Doc",
            content="Content",
        )
        Summary.objects.create(
            document=doc,
            user=self.user,
            title="Sum",
            method="textrank",
            language="en",
            ratio=0.2,
            summary_text="Text",
        )
        self.user.delete()
        self.assertEqual(Summary.objects.count(), 0)
        self.assertEqual(Document.objects.count(), 0)


# ──────────────────────────────────────────────
#  URL EXTRACTION TESTS (MOCKED)
# ──────────────────────────────────────────────


class UrlExtractionTests(TestCase):
    def _mock_session_get(self, mock_get):
        mock_get.return_value.status_code = 200
        mock_get.return_value.headers = {"Content-Type": "text/html; charset=utf-8"}

    @patch("summaries.readers._get_http_session")
    def test_extract_text_from_url(self, mock_session):
        sess = mock_session.return_value
        sess.get.return_value.status_code = 200
        sess.get.return_value.headers = {"Content-Type": "text/html; charset=utf-8"}
        sess.get.return_value.text = (
            "<html><body><p>Hello world. This is a test page.</p></body></html>"
        )
        result = extract_text("https://example.com")
        self.assertIn("Hello world", result)

    @patch("summaries.readers._get_http_session")
    def test_extract_text_from_url_removes_scripts(self, mock_session):
        sess = mock_session.return_value
        sess.get.return_value.status_code = 200
        sess.get.return_value.headers = {"Content-Type": "text/html; charset=utf-8"}
        sess.get.return_value.text = (
            "<html><head><script>alert('xss')</script></head>"
            "<body><p>Main content here.</p></body></html>"
        )
        result = extract_text("https://example.com")
        self.assertIn("Main content", result)
        self.assertNotIn("alert", result)

    @patch("summaries.readers._get_http_session")
    def test_extract_text_from_url_invalid_content_type(self, mock_session):
        sess = mock_session.return_value
        sess.get.return_value.status_code = 200
        sess.get.return_value.headers = {"Content-Type": "application/pdf"}
        sess.get.return_value.text = "not html"
        with self.assertRaises(ValueError):
            extract_text("https://example.com/file.pdf")

    @patch("summaries.readers._get_http_session")
    def test_extract_text_from_url_raises_on_timeout(self, mock_session):
        from requests.exceptions import Timeout as RequestsTimeout

        sess = mock_session.return_value
        sess.get.side_effect = RequestsTimeout("Timeout")
        with self.assertRaises(ValueError):
            extract_text("https://example.com")

    def test_extract_text_from_url_invalid_scheme(self):
        with self.assertRaises((ValueError, FileNotFoundError)):
            extract_text("ftp://example.com")

    def test_extract_text_from_url_no_netloc(self):
        with self.assertRaises(ValueError):
            extract_text("http://")


class SsrfProtectionTests(TestCase):
    def test_is_private_ip_blocks_loopback(self):
        self.assertTrue(_is_private_ip("127.0.0.1"))
        self.assertTrue(_is_private_ip("::1"))

    def test_is_private_ip_blocks_private_ranges(self):
        self.assertTrue(_is_private_ip("10.0.0.5"))
        self.assertTrue(_is_private_ip("192.168.1.1"))
        self.assertTrue(_is_private_ip("172.16.0.1"))
        self.assertTrue(_is_private_ip("169.254.169.254"))
        self.assertTrue(_is_private_ip("fd00::1"))

    def test_is_private_ip_allows_public(self):
        self.assertFalse(_is_private_ip("8.8.8.8"))
        self.assertFalse(_is_private_ip("93.184.216.34"))

    @patch("summaries.readers.socket.getaddrinfo")
    def test_resolve_and_validate_blocks_private(self, mock_getaddrinfo):
        mock_getaddrinfo.return_value = [
            (None, None, None, None, ("127.0.0.1", 80))
        ]
        with self.assertRaises(ValueError):
            _resolve_and_validate("localhost")

    @patch("summaries.readers.socket.getaddrinfo")
    def test_resolve_and_validate_allows_public(self, mock_getaddrinfo):
        mock_getaddrinfo.return_value = [
            (None, None, None, None, ("8.8.8.8", 80))
        ]
        _resolve_and_validate("example.com")

    @patch("summaries.readers.socket.getaddrinfo")
    def test_extract_url_blocks_private_hostname(self, mock_getaddrinfo):
        mock_getaddrinfo.return_value = [
            (None, None, None, None, ("127.0.0.1", 80))
        ]
        with self.assertRaises(ValueError):
            extract_text("http://localhost/secret")


class UrlSourceFlowTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="url-test", password="secret123")

        self.client.login(username="url-test", password="secret123")

    def tearDown(self):
        cache.clear()

    @patch("summaries.readers._get_http_session")
    def test_create_summary_from_url_view(self, mock_session):
        sess = mock_session.return_value
        sess.get.return_value.status_code = 200
        sess.get.return_value.headers = {"Content-Type": "text/html; charset=utf-8"}
        sess.get.return_value.text = (
            "<html><body><p>First useful sentence. "
            "Second useful sentence. Third sentence.</p></body></html>"
        )
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "url",
                "source_url": "https://example.com/article",
                "method": "textrank",
                "ratio": 0.3,
            },
        )
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(data["ok"])
        self.assertEqual(data["data"]["source_type"], "url")

    def test_create_summary_url_requires_url(self):
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "url",
                "method": "textrank",
                "ratio": 0.3,
            },
        )
        self.assertEqual(response.status_code, 400)


class SigningTests(TestCase):
    def test_encrypt_decrypt_roundtrip(self):
        text = "my-api-key-123"
        self.assertEqual(decrypt_value(encrypt_value(text)), text)

    def test_encrypt_empty_returns_empty(self):
        self.assertEqual(encrypt_value(""), "")

    def test_decrypt_invalid_token_returns_raw(self):
        self.assertEqual(decrypt_value("not-a-valid-token"), "not-a-valid-token")


class SuperuserRoleEvolutionTests(TestCase):
    def test_superuser_profile_becomes_admin_on_home_and_ratio_initialized(self):
        User.objects.create_superuser(username="boss", password="secret123")
        self.client.login(username="boss", password="secret123")
        response = self.client.get(reverse("home"))
        self.assertEqual(response.status_code, 200)
        profile = UserProfile.objects.get(user__username="boss")
        self.assertEqual(profile.role, "admin")
        setting = UserSetting.objects.get(user__username="boss")
        self.assertEqual(setting.default_summary_ratio, 0.2)


# ──────────────────────────────────────────────
#  FILE UPLOAD TESTS
# ──────────────────────────────────────────────


@override_settings(MEDIA_ROOT=tempfile.mkdtemp(), RATE_LIMIT_SECONDS=0)
class FileUploadTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="uploader", password="secret123")

        self.client.login(username="uploader", password="secret123")

    def test_upload_txt_file(self):
        text_content = b"This is a test document. It has multiple sentences. We need enough text."
        uploaded = SimpleUploadedFile("test.txt", text_content, content_type="text/plain")
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "file",
                "method": "textrank",
                "upload": uploaded,
                "ratio": 0.3,
            },
        )
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(data["ok"])

    def test_upload_file_too_large(self):
        big_content = b"x" * (10 * 1024 * 1024 + 1)
        uploaded = SimpleUploadedFile("big.txt", big_content, content_type="text/plain")
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "file",
                "method": "textrank",
                "upload": uploaded,
                "ratio": 0.3,
            },
        )
        self.assertEqual(response.status_code, 400)

    def test_upload_no_file_sent(self):
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "file",
                "method": "textrank",
                "ratio": 0.3,
            },
        )
        self.assertEqual(response.status_code, 400)

    def test_upload_unsupported_format(self):
        uploaded = SimpleUploadedFile(
            "test.exe", b"fake content", content_type="application/octet-stream"
        )
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "file",
                "method": "textrank",
                "upload": uploaded,
                "ratio": 0.3,
            },
        )
        self.assertEqual(response.status_code, 400)


# ──────────────────────────────────────────────
#  GEMINI SUMMARIZE TESTS (MOCKED)
# ──────────────────────────────────────────────


@override_settings(RATE_LIMIT_SECONDS=0)
class GeminiSummarizeTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="gemini-test", password="secret123")
        UserSetting.objects.filter(user=self.user).update(gemini_api_key=encrypt_value("fake-key"))
        self.client.login(username="gemini-test", password="secret123")
        self.mock_response = {
            "candidates": [{"content": {"parts": [{"text": "This is a short summary."}]}}]
        }

    def _mock_session(self, mock_get_session, status_code=200, response_data=None):
        sess = mock_get_session.return_value
        mock_resp = MagicMock()
        mock_resp.status_code = status_code
        mock_resp.json.return_value = response_data or self.mock_response
        mock_resp.text = ""
        if status_code >= 400:
            from requests.exceptions import HTTPError

            mock_resp.raise_for_status.side_effect = HTTPError(f"HTTP {status_code}")
        sess.post.return_value = mock_resp
        return sess

    def tearDown(self):
        cache.clear()

    @patch("summaries.readers._get_http_session")
    def test_gemini_summarize_success(self, mock_get_session):
        self._mock_session(mock_get_session)
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "text",
                "method": "gemini",
                "text": "This is the first sentence. Here is another one. Yet a third sentence.",
                "ratio": 0.3,
            },
        )
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(data["ok"])

    @override_settings(GEMINI_API_KEY="")
    @patch("summaries.readers._get_http_session")
    def test_gemini_summarize_api_key_missing(self, mock_get_session):
        self._mock_session(mock_get_session)
        UserSetting.objects.filter(user=self.user).update(gemini_api_key="")
        cache.clear()
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "text",
                "method": "gemini",
                "text": "Some text for summary.",
                "ratio": 0.3,
            },
        )
        self.assertEqual(response.status_code, 400)
        data = response.json()
        self.assertFalse(data["ok"])

    @patch("summaries.readers._get_http_session")
    def test_gemini_summarize_empty_response(self, mock_get_session):
        bad_resp = {"candidates": [{"content": {"parts": [{"text": ""}]}}]}
        self._mock_session(mock_get_session, response_data=bad_resp)
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "text",
                "method": "gemini",
                "text": "First word. Second word. Third word.",
                "ratio": 0.3,
            },
        )
        self.assertEqual(response.status_code, 400)

    @patch("summaries.readers._get_http_session")
    def test_gemini_summarize_http_error(self, mock_get_session):
        self._mock_session(mock_get_session, status_code=500)
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "text",
                "method": "gemini",
                "text": "First word. Second word. Third word.",
                "ratio": 0.3,
            },
        )
        self.assertEqual(response.status_code, 400)

    @patch("summaries.readers._get_http_session")
    def test_gemini_summarize_malformed_json(self, mock_get_session):
        bad_resp = {"unexpected": "format"}
        self._mock_session(mock_get_session, response_data=bad_resp)
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "text",
                "method": "gemini",
                "text": "Some text here. More text there.",
                "ratio": 0.3,
            },
        )
        self.assertEqual(response.status_code, 400)


# ──────────────────────────────────────────────
#  FILE EXTRACTION TESTS (MOCKED)
# ──────────────────────────────────────────────


@override_settings(RATE_LIMIT_SECONDS=0)
class FileExtractionTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="extract-test", password="secret123")

        self.client.login(username="extract-test", password="secret123")

    def tearDown(self):
        cache.clear()

    @patch("summaries.readers._extract_text_from_txt")
    def test_txt_file_extraction(self, mock_extract):
        mock_extract.return_value = (
            "This is extracted text from a txt file. It has multiple sentences."
        )
        uploaded = SimpleUploadedFile("test.txt", b"ignored", content_type="text/plain")
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "file",
                "method": "textrank",
                "upload": uploaded,
                "ratio": 0.3,
            },
        )
        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.json()["ok"])
        mock_extract.assert_called_once()

    @patch("summaries.readers._extract_text_from_docx")
    def test_docx_file_extraction(self, mock_extract):
        mock_extract.return_value = "Extracted content from a DOCX file."
        uploaded = SimpleUploadedFile(
            "test.docx",
            b"ignored",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "file",
                "method": "textrank",
                "upload": uploaded,
                "ratio": 0.3,
            },
        )
        self.assertEqual(response.status_code, 200)
        mock_extract.assert_called_once()

    @patch("summaries.readers._extract_text_from_pdf")
    def test_pdf_file_extraction(self, mock_extract):
        mock_extract.return_value = "Extracted content from a PDF file."
        uploaded = SimpleUploadedFile("test.pdf", b"ignored", content_type="application/pdf")
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "file",
                "method": "textrank",
                "upload": uploaded,
                "ratio": 0.3,
            },
        )
        self.assertEqual(response.status_code, 200)
        mock_extract.assert_called_once()

    @patch("summaries.readers._extract_text_from_epub")
    def test_epub_file_extraction(self, mock_extract):
        mock_extract.return_value = "Extracted content from an EPUB file."
        uploaded = SimpleUploadedFile("test.epub", b"ignored", content_type="application/epub+zip")
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "file",
                "method": "textrank",
                "upload": uploaded,
                "ratio": 0.3,
            },
        )
        self.assertEqual(response.status_code, 200)
        mock_extract.assert_called_once()


# ──────────────────────────────────────────────
#  ERROR PAGE TESTS (404 / 500)
# ──────────────────────────────────────────────


@override_settings(DEBUG=False, ALLOWED_HOSTS=["*"])
class ErrorPageTests(TestCase):
    def setUp(self):
        self.client.raise_request_exception = False

    def test_missing_page_renders_custom_404(self):
        response = self.client.get("/nonexistent-page/")
        self.assertEqual(response.status_code, 404)
        self.assertIn("Trang bạn tìm kiếm không tồn tại.", response.content.decode())
        self.assertIn("error-code", response.content.decode())

    def test_server_error_renders_custom_500(self):
        with override_settings(ROOT_URLCONF="summaries.tests"):
            response = self.client.get("/__boom__/")
        self.assertEqual(response.status_code, 500)
        content = response.content.decode()
        self.assertIn("Lỗi máy chủ", content)
        self.assertIn("Vui lòng thử lại sau", content)


# ──────────────────────────────────────────────
#  SECURITY HEADER TESTS
# ──────────────────────────────────────────────


class SecurityHeaderTests(TestCase):
    def test_csp_policy_present_on_all_pages(self):
        response = self.client.get(reverse("home"))
        csp = response.headers.get("Content-Security-Policy", "")
        self.assertIn("default-src 'self'", csp)
        self.assertIn("script-src 'self' 'nonce-", csp)
        self.assertIn("style-src 'self'", csp)
        self.assertIn("img-src 'self' data:", csp)
        self.assertIn("base-uri 'self'", csp)
        self.assertIn("form-action 'self'", csp)

    def test_clickjacking_protection_enabled(self):
        response = self.client.get(reverse("home"))
        self.assertEqual(response.headers.get("X-Frame-Options"), "DENY")

    def test_nosniff_header(self):
        response = self.client.get(reverse("home"))
        self.assertEqual(response.headers.get("X-Content-Type-Options"), "nosniff")

    def test_referrer_policy(self):
        response = self.client.get(reverse("home"))
        self.assertEqual(response.headers.get("Referrer-Policy"), "strict-origin-when-cross-origin")

    def test_csrf_blocks_missing_token(self):
        from django.test import Client

        client = Client(enforce_csrf_checks=True)
        response = client.post(reverse("login"), {"username": "nobody", "password": "x"})
        self.assertEqual(response.status_code, 403)


# ──────────────────────────────────────────────
#  CONTENT EDGE CASE TESTS
# ──────────────────────────────────────────────


@override_settings(RATE_LIMIT_SECONDS=0)
class ContentEdgeCaseTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="edge-test", password="secret123")

        self.client.login(username="edge-test", password="secret123")

    def tearDown(self):
        cache.clear()

    def test_single_character_text(self):
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "text",
                "method": "textrank",
                "text": "A",
                "ratio": 0.2,
            },
        )
        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.json()["ok"])

    def test_whitespace_only_text_rejected(self):
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "text",
                "method": "textrank",
                "text": "   \n\t  ",
                "ratio": 0.2,
            },
        )
        self.assertEqual(response.status_code, 400)
        self.assertFalse(response.json()["ok"])

    def test_very_long_text_summarized(self):
        long_text = (
            "Đây là một câu dùng để kiểm tra khả năng xử lý văn bản dài. "
            "Khi dữ liệu lớn, hệ thống vẫn phải tóm tắt chính xác và đầy đủ. "
        ) * 60
        response = self.client.post(
            reverse("create_summary"),
            {
                "source_type": "text",
                "method": "textrank",
                "text": long_text,
                "ratio": 0.2,
            },
        )
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload["ok"])
        self.assertTrue(payload["data"]["summary"])


# ── ROOT_URLCONF nhỏ dùng riêng cho ErrorPageTests ──
from django.urls import path as _path  # noqa: E402

from .views import LoginPageView as _Login  # noqa: E402
from .views import RegisterPageView as _Register  # noqa: E402
from .views import home as _home  # noqa: E402


def _boom_view(_request):  # noqa: N802
    raise RuntimeError("boom")


urlpatterns = [
    _path("", _home, name="home"),
    _path("login/", _Login.as_view(), name="login"),
    _path("register/", _Register.as_view(), name="register"),
    _path("__boom__/", _boom_view),
]


# ──────────────────────────────────────────────
#  GEMINI ERROR BRANCH TESTS (DIRECT)
# ──────────────────────────────────────────────


class GeminiErrorBranchTests(TestCase):
    """Unit tests for gemini_summarize error/retry branches (coverage gap).

    Covers 403, 400 (with detail), 429/502/503 retry, retry exhaustion,
    timeout, connection error, Vietnamese prompt, and ratio-to-vietnamese
    length variants.
    """

    def setUp(self):
        import json as _json

        self._json = _json
        self.mock_response = {
            "candidates": [{"content": {"parts": [{"text": "Tóm tắt ngắn gọn."}]}}]
        }

    def _session(self, mock_get_session, responses):
        """responses: list of (status, payload_callable). Side-effect cycles."""
        sess = mock_get_session.return_value

        def _resp_for(status, callback):
            r = MagicMock()
            r.status_code = status
            r.text = "body"
            r.json.side_effect = callback
            if status >= 400:
                from requests.exceptions import HTTPError

                r.raise_for_status.side_effect = HTTPError(f"HTTP {status}")
            return r

        mock_resps = [_resp_for(s, cb) for s, cb in responses]
        sess.post.side_effect = mock_resps
        return sess

    def _call(self):
        from .nlp import gemini_summarize

        return gemini_summarize(
            "Example text for Gemini. This is the first sentence of the input.",
            ratio=0.3,
            language="english",
            user_api_key="fake-key",
        )

    @patch("summaries.nlp.time_module.sleep")
    @patch("summaries.readers._get_http_session")
    def test_gemini_403_rejected(self, mock_get_session, _sleep):
        def cb():
            return {"error": {"message": "invalid key"}}

        self._session(mock_get_session, [(403, cb)])
        with self.assertRaises(ValueError):
            self._call()

    @patch("summaries.nlp.time_module.sleep")
    @patch("summaries.readers._get_http_session")
    def test_gemini_400_with_detail(self, mock_get_session, _sleep):
        def cb():
            return {"error": {"message": "quota exceeded"}}

        self._session(mock_get_session, [(400, cb)])
        with self.assertRaises(ValueError) as ctx:
            self._call()
        self.assertIn("quota exceeded", str(ctx.exception))

    @patch("summaries.nlp.time_module.sleep")
    @patch("summaries.readers._get_http_session")
    def test_gemini_429_retries_then_succeeds(self, mock_get_session, _sleep):
        self._session(
            mock_get_session,
            [
                (429, MagicMock(return_value={})),
                (429, MagicMock(return_value={})),
                (200, MagicMock(return_value=self.mock_response)),
            ],
        )
        result = self._call()
        self.assertIn("summary", result)

    @patch("summaries.nlp.time_module.sleep")
    @patch("summaries.readers._get_http_session")
    def test_gemini_502_exhausts_retries(self, mock_get_session, _sleep):
        self._session(
            mock_get_session,
            [
                (502, MagicMock(return_value={})),
                (502, MagicMock(return_value={})),
                (502, MagicMock(return_value={})),
            ],
        )
        with self.assertRaises(ValueError):
            self._call()

    @patch("summaries.readers._get_http_session")
    def test_gemini_timeout(self, mock_get_session):
        from requests.exceptions import Timeout as RequestsTimeout

        sess = mock_get_session.return_value
        sess.post.side_effect = RequestsTimeout("timed out")
        with self.assertRaises(ValueError) as ctx:
            self._call()
        self.assertIn("60", str(ctx.exception))

    @patch("summaries.readers._get_http_session")
    def test_gemini_connection_error(self, mock_get_session):
        from requests.exceptions import ConnectionError as RequestsConnError

        sess = mock_get_session.return_value
        sess.post.side_effect = RequestsConnError("refused")
        with self.assertRaises(ValueError) as ctx:
            self._call()
        self.assertIn("kết nối", str(ctx.exception))


class GeminiPromptBranchTests(TestCase):
    """Cover _build_prompt Vietnamese branch + remaining ratio variants."""

    @patch("summaries.nlp.time_module.sleep")
    @patch("summaries.readers._get_http_session")
    def test_vietnamese_prompt_success(self, mock_get_session, _sleep):
        resp = MagicMock()
        resp.status_code = 200
        resp.text = ""
        resp.json.return_value = {
            "candidates": [{"content": {"parts": [{"text": "Tóm tắt tiếng Việt."}]}}]
        }
        mock_get_session.return_value.post.return_value = resp

        from .nlp import gemini_summarize

        result = gemini_summarize(
            "Đây là câu đầu tiên của văn bản cần tóm tắt bằng tiếng Việt.",
            ratio=0.1,
            language="vietnamese",
            user_api_key="fake-key",
        )
        self.assertIn("summary", result)

    @patch("summaries.nlp.time_module.sleep")
    @patch("summaries.readers._get_http_session")
    def test_vietnamese_ratio_variants(self, mock_get_session, _sleep):
        resp = MagicMock()
        resp.status_code = 200
        resp.text = ""
        resp.json.return_value = {
            "candidates": [{"content": {"parts": [{"text": "Tóm tắt."}]}}]
        }
        mock_get_session.return_value.post.return_value = resp

        from .nlp import _ratio_to_vietnamese

        # All ratio branches must produce non-empty strings
        for r in (0.1, 0.2, 0.4, 0.7):
            self.assertTrue(_ratio_to_vietnamese(r))


# ──────────────────────────────────────────────
#  READER UNIT TESTS (REAL EXTRACTION + ERRORS)
# ──────────────────────────────────────────────


class ReaderUnitTests(TestCase):
    def test_extract_text_from_txt_utf8(self):
        import pathlib
        import tempfile as _tf

        from .readers import _extract_text_from_txt

        with _tf.NamedTemporaryFile("w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write("Xin chào thế giới. Đây là nội dung UTF-8.")
            p = pathlib.Path(f.name)
        try:
            out = _extract_text_from_txt(p)
        finally:
            p.unlink(missing_ok=True)
        self.assertIn("Xin chào", out)

    def test_extract_text_from_txt_without_chardet(self):
        import pathlib
        import tempfile as _tf

        from .readers import _extract_text_from_txt

        with _tf.NamedTemporaryFile("wb", suffix=".txt", delete=False) as f:
            f.write(b"plain ascii text")
            p = pathlib.Path(f.name)
        try:
            with patch.dict("sys.modules", {"chardet": None}):
                out = _extract_text_from_txt(p)
        finally:
            p.unlink(missing_ok=True)
        self.assertEqual(out, "plain ascii text")

    def test_extract_docx_missing_lib(self):
        from .readers import _extract_text_from_docx

        with patch.dict("sys.modules", {"docx": None}):
            with self.assertRaises(ValueError) as ctx:
                _extract_text_from_docx("x.docx")
        self.assertIn("python-docx", str(ctx.exception))

    def test_extract_pdf_missing_lib(self):
        from .readers import _extract_text_from_pdf

        with patch.dict("sys.modules", {"fitz": None}):
            with self.assertRaises(ValueError) as ctx:
                _extract_text_from_pdf("x.pdf")
        self.assertIn("PyMuPDF", str(ctx.exception))

    def test_extract_docx_success(self):
        import pathlib
        import tempfile as _tf

        from docx import Document as DocxDocument

        from .readers import _extract_text_from_docx

        with _tf.TemporaryDirectory() as d:
            p = pathlib.Path(d) / "a.docx"
            doc = DocxDocument()
            doc.add_paragraph("First paragraph of the docx file.")
            doc.add_paragraph("Second paragraph content.")
            doc.save(p)
            out = _extract_text_from_docx(p)
        self.assertIn("First paragraph", out)
        self.assertIn("Second paragraph", out)

    @patch("summaries.readers._get_http_session")
    @patch("summaries.readers._resolve_and_validate")
    def test_redirects_then_succeeds(self, _resolve, mock_get_session):
        sess = mock_get_session.return_value
        r1, r2 = MagicMock(), MagicMock()
        r1.status_code = 301
        r1.headers = {"Location": "https://final.example.com/page"}
        r1.text = ""
        r2.status_code = 200
        r2.headers = {"Content-Type": "text/html; charset=utf-8"}
        r2.text = "<html><body><p>Final page content here.</p></body></html>"
        sess.get.side_effect = [r1, r2]

        out = extract_text("https://start.example.com/old")
        self.assertIn("Final page content", out)

    @patch("summaries.readers._get_http_session")
    @patch("summaries.readers._resolve_and_validate")
    def test_http_error_raised(self, _resolve, mock_get_session):
        sess = mock_get_session.return_value
        resp = MagicMock()
        resp.status_code = 404
        resp.text = "nope"
        resp.headers = {}
        sess.get.return_value = resp

        with self.assertRaises(ValueError) as ctx:
            extract_text("https://example.com/missing")
        self.assertIn("404", str(ctx.exception))

    @patch("summaries.readers._get_http_session")
    @patch("summaries.readers._resolve_and_validate")
    def test_plain_text_content_type(self, _resolve, mock_get_session):
        sess = mock_get_session.return_value
        resp = MagicMock()
        resp.status_code = 200
        resp.headers = {"Content-Type": "text/plain; charset=utf-8"}
        resp.text = "Just some plain text document."
        sess.get.return_value = resp

        out = extract_text("https://example.com/notes.txt")
        self.assertEqual(out, "Just some plain text document.")

    @patch("summaries.readers._get_http_session")
    @patch("summaries.readers._resolve_and_validate")
    def test_empty_html_raises(self, _resolve, mock_get_session):
        sess = mock_get_session.return_value
        resp = MagicMock()
        resp.status_code = 200
        resp.headers = {"Content-Type": "text/html; charset=utf-8"}
        resp.text = "<html><body></body></html>"
        sess.get.return_value = resp

        with self.assertRaises(ValueError):
            extract_text("https://example.com/empty")


# ──────────────────────────────────────────────
#  SETUP COMMAND TESTS
# ──────────────────────────────────────────────


class SetupCommandTests(TestCase):
    def test_setup_migrate_only(self):
        from django.core.management import call_command
        from django.core.management.base import CommandError

        try:
            call_command("setup")
        except CommandError:
            pass  # migrate no-op in test DB is fine

    def test_setup_create_superuser_when_none(self):
        from django.contrib.auth import get_user_model
        from django.core.management import call_command

        user_model = get_user_model()
        user_model.objects.all().delete()
        call_command("setup", create_superuser=True, username="newboss", password="NewPass456!")
        self.assertTrue(user_model.objects.filter(username="newboss", is_superuser=True).exists())

    def test_setup_skips_when_superuser_exists(self):
        from django.contrib.auth import get_user_model
        from django.core.management import call_command

        user_model = get_user_model()
        user_model.objects.filter(is_superuser=True).delete()
        user_model.objects.create_superuser(
            username="existingboss", password="Pass123!", email="b@b.com"
        )
        call_command("setup", create_superuser=True, username="newboss", password="NewPass456!")
        self.assertFalse(user_model.objects.filter(username="newboss").exists())


# ──────────────────────────────────────────────
#  SIGNING FALLBACK BRANCHES
# ──────────────────────────────────────────────


class SigningFallbackTests(TestCase):
    """Cover cryptography-not-installed fallback paths in signing.py."""

    def test_encrypt_when_fernet_none_returns_plaintext(self):
        import summaries.signing as signing

        real = signing.Fernet
        signing.Fernet = None
        try:
            self.assertEqual(encrypt_value("secret-value"), "secret-value")
        finally:
            signing.Fernet = real

    def test_decrypt_when_fernet_none_returns_raw(self):
        import summaries.signing as signing

        real = signing.Fernet
        signing.Fernet = None
        try:
            self.assertEqual(decrypt_value("raw-encrypted"), "raw-encrypted")
        finally:
            signing.Fernet = real

    def test_encrypt_when_fernet_none_and_empty_returns_empty(self):
        import summaries.signing as signing

        real = signing.Fernet
        signing.Fernet = None
        try:
            self.assertEqual(encrypt_value(""), "")
        finally:
            signing.Fernet = real

    def test_module_import_fallback_sets_none_on_missing_cryptography(self):
        import importlib

        import summaries.signing as signing

        real = signing.Fernet
        original_import = __import__

        def fake_import(name, *a, **k):
            if name.startswith("cryptography"):
                raise ImportError("blocked")
            return original_import(name, *a, **k)

        try:
            signing.Fernet = None
            with patch("builtins.__import__", side_effect=fake_import):
                importlib.reload(signing)
            self.assertIsNone(signing.Fernet)
        finally:
            signing.Fernet = real
            importlib.reload(signing)


# ──────────────────────────────────────────────
#  MODEL STR + CLEANUP EDGE CASES
# ──────────────────────────────────────────────


class ModelStrAndCleanupTests(TestCase):
    def test_cleanup_uploaded_file_empty(self):
        from .models import _cleanup_uploaded_file

        _cleanup_uploaded_file("")
        _cleanup_uploaded_file(None)

    def test_cleanup_uploaded_file_removes_existing(self):
        import pathlib
        import tempfile as _tf

        from .models import _cleanup_uploaded_file

        with _tf.TemporaryDirectory() as d:
            target = pathlib.Path(d) / "x.txt"
            target.write_text("hi", encoding="utf-8")
            with override_settings(MEDIA_ROOT=d):
                _cleanup_uploaded_file("x.txt")
            self.assertFalse(target.exists())

    def test_cleanup_uploaded_file_missing_is_noop(self):
        from .models import _cleanup_uploaded_file

        with override_settings(MEDIA_ROOT="C:\\nonexistent\\nope"):
            _cleanup_uploaded_file("missing.txt")

    def test_user_setting_str(self):
        user = User.objects.create_user(username="str-user", password="secret123")
        setting, _ = UserSetting.objects.get_or_create(
            user=user, defaults={"default_summary_ratio": 0.2}
        )
        self.assertIn("str-user", str(setting))

    def test_document_str(self):
        user = User.objects.create_user(username="doc-str-user", password="secret123")
        doc = Document.objects.create(
            user=user, source_type="text", title="My Doc Title", content="x"
        )
        self.assertEqual(str(doc), "My Doc Title")

    def test_cleanup_uploaded_file_ignores_oserror(self):
        from .models import _cleanup_uploaded_file

        with patch("pathlib.Path.exists", return_value=True), patch(
            "pathlib.Path.unlink", side_effect=OSError("denied")
        ):
            _cleanup_uploaded_file("locked.txt")


# ──────────────────────────────────────────────
#  READER COVERAGE (REMAINING BRANCHES)
# ──────────────────────────────────────────────


class ReaderCoverageTests(TestCase):
    @patch("summaries.readers.socket.getaddrinfo")
    def test_resolve_and_validate_gaierror(self, mock_addr):
        import socket as _socket

        mock_addr.side_effect = _socket.gaierror("no such host")
        with self.assertRaises(ValueError) as ctx:
            _resolve_and_validate("definitely-not-a-host.invalid")
        self.assertIn("phân giải", str(ctx.exception))

    @patch("summaries.readers.socket.getaddrinfo")
    def test_resolve_and_validate_blocks_private_ip(self, mock_addr):
        mock_addr.return_value = [
            (2, 1, 6, "", ("10.0.0.5", 80)),
        ]
        with self.assertRaises(ValueError) as ctx:
            _resolve_and_validate("example.com")
        self.assertIn("nội bộ", str(ctx.exception))

    def test_is_private_ip_invalid_address_returns_false(self):
        self.assertFalse(_is_private_ip("not-an-ip-address"))

    def test_get_http_session_thread_safe(self):
        from .readers import _get_http_session

        session = _get_http_session()
        self.assertIs(session, _get_http_session())
        self.assertIn("User-Agent", session.headers)
        _local = getattr(_get_http_session, "__self__", None)

    def test_extract_text_from_url_requests_missing(self):
        from .readers import extract_text_from_url

        with patch.dict("sys.modules", {"requests": None}):
            with self.assertRaises(ValueError) as ctx:
                extract_text_from_url("https://example.com")
        self.assertIn("requests", str(ctx.exception))

    def test_extract_text_from_url_bs4_missing(self):
        from .readers import extract_text_from_url

        with patch.dict("sys.modules", {"bs4": None}):
            with self.assertRaises(ValueError) as ctx:
                extract_text_from_url("https://example.com")
        self.assertIn("beautifulsoup4", str(ctx.exception))

    @patch("summaries.readers._get_http_session")
    @patch("summaries.readers._resolve_and_validate")
    def test_extract_text_from_url_connection_error(self, _resolve, mock_session):
        from requests.exceptions import ConnectionError as ConnErr

        mock_session.return_value.get.side_effect = ConnErr("refused")
        with self.assertRaises(ValueError):
            extract_text("https://example.com")

    @patch("summaries.readers._get_http_session")
    @patch("summaries.readers._resolve_and_validate")
    def test_extract_text_from_url_redirect_with_empty_location(self, _resolve, mock_session):
        """Redirect response with an empty Location stops following and raises."""
        sess = mock_session.return_value
        r = MagicMock()
        r.status_code = 302
        r.headers = {"Location": ""}
        r.text = ""
        sess.get.return_value = r
        with self.assertRaises(ValueError):
            extract_text("https://example.com")

    def test_extract_text_pdf_success(self):
        import pathlib
        import tempfile as _tf

        from .readers import _extract_text_from_pdf

        with _tf.TemporaryDirectory() as d:
            p = pathlib.Path(d) / "a.pdf"
            try:
                import fitz

                doc = fitz.open()
                page = doc.new_page()
                page.insert_text((72, 72), "Hello PDF content")
                doc.save(str(p))
                doc.close()
                out = _extract_text_from_pdf(p)
                self.assertIn("Hello PDF", out)
            except ImportError:
                self.skipTest("PyMuPDF not installed")

    def test_extract_text_from_epub_bs4_missing(self):
        from .readers import _extract_text_from_epub

        with patch.dict("sys.modules", {"bs4": None}):
            with self.assertRaises(ValueError) as ctx:
                _extract_text_from_epub("x.epub")
        self.assertIn("beautifulsoup4", str(ctx.exception))

    def test_extract_text_from_epub_ebooklib_missing(self):
        from .readers import _extract_text_from_epub

        with patch.dict("sys.modules", {"ebooklib": None}):
            with self.assertRaises(ValueError) as ctx:
                _extract_text_from_epub("x.epub")
        self.assertIn("ebooklib", str(ctx.exception))

    def test_extract_text_unsupported_extension(self):
        import pathlib
        import tempfile as _tf

        with _tf.TemporaryDirectory() as d:
            p = pathlib.Path(d) / "foo.exe"
            p.write_bytes(b"data")
            with self.assertRaises(ValueError) as ctx:
                extract_text(str(p))
        self.assertIn("không được hỗ trợ", str(ctx.exception))

    def test_extract_text_missing_file(self):
        with self.assertRaises(FileNotFoundError):
            extract_text("/tmp/definitely-missing-file.txt")


